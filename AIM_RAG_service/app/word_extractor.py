"""Word (.docx / .doc) text + embedded-image extraction for order upload.

.docx  → python-docx text (paragraphs + tables) + ZIP media images
.doc   → LibreOffice convert to .docx, then same path
         (misnamed .docx saved as .doc detected via ZIP/PK header)
"""

from __future__ import annotations

import io
import os
import shutil
import subprocess
import tempfile
import zipfile
from typing import List, Optional, Tuple

_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff", ".gif"}


def _is_zip_docx_bytes(data: bytes) -> bool:
    """True when bytes are a ZIP package (real .docx / OOXML)."""
    return bool(data) and data[:2] == b"PK"


def _find_soffice() -> Optional[str]:
    for name in ("soffice", "soffice.exe", "libreoffice"):
        path = shutil.which(name)
        if path:
            return path
    for candidate in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/usr/bin/soffice",
        "/usr/bin/libreoffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ):
        if os.path.isfile(candidate):
            return candidate
    return None


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract plain text from a .docx (paragraphs + table cells). May be empty."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise ValueError(
            "python-docx is not installed. Add python-docx to requirements and reinstall."
        ) from exc

    document = DocxDocument(io.BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join((cell.text or "").split())
                if cell_text:
                    cells.append(cell_text)
            if cells:
                deduped = []
                for cell_text in cells:
                    if not deduped or deduped[-1] != cell_text:
                        deduped.append(cell_text)
                parts.append(" | ".join(deduped))

    return "\n".join(parts).strip()


def extract_images_from_docx_bytes(file_bytes: bytes) -> List[Tuple[str, bytes]]:
    """
    Pull embedded images from a .docx ZIP (word/media/*).

    Returns list of (filename, image_bytes), largest first (order screenshots
    are usually the biggest media files).
    """
    if not _is_zip_docx_bytes(file_bytes):
        return []

    images: List[Tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
            for name in archive.namelist():
                normalized = name.replace("\\", "/")
                if not normalized.lower().startswith("word/media/"):
                    continue
                ext = os.path.splitext(normalized)[1].lower()
                if ext not in _IMAGE_EXTENSIONS:
                    continue
                try:
                    data = archive.read(name)
                except Exception:
                    continue
                # Skip tiny icons / bullets; keep order screenshots (usually much larger)
                if not data or len(data) < 1_500:
                    continue
                images.append((os.path.basename(normalized), data))
    except zipfile.BadZipFile:
        return []

    images.sort(key=lambda item: len(item[1]), reverse=True)
    return images


def _convert_doc_to_docx_bytes(file_bytes: bytes) -> bytes:
    """Convert legacy .doc to .docx using LibreOffice (when installed)."""
    soffice = _find_soffice()
    if not soffice:
        raise ValueError(
            "Legacy .doc files require LibreOffice for conversion. "
            "Install LibreOffice, or save/export the order as .docx or .pdf and re-upload."
        )

    with tempfile.TemporaryDirectory(prefix="word_doc_") as tmp_dir:
        input_path = os.path.join(tmp_dir, "input.doc")
        with open(input_path, "wb") as handle:
            handle.write(file_bytes)

        try:
            completed = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--nologo",
                    "--nolockcheck",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    tmp_dir,
                    input_path,
                ],
                check=False,
                capture_output=True,
                timeout=120,
            )
        except subprocess.TimeoutExpired as exc:
            raise ValueError("LibreOffice timed out while converting .doc to .docx.") from exc

        docx_files = [
            name for name in os.listdir(tmp_dir) if name.lower().endswith(".docx")
        ]
        if not docx_files:
            stderr = (completed.stderr or b"").decode("utf-8", errors="replace").strip()
            detail = f" LibreOffice: {stderr}" if stderr else ""
            raise ValueError(
                f"Failed to convert legacy .doc to .docx.{detail} "
                "Save/export as .docx or .pdf and re-upload."
            )

        output_path = os.path.join(tmp_dir, docx_files[0])
        with open(output_path, "rb") as handle:
            return handle.read()


def normalize_word_to_docx_bytes(file_bytes: bytes, filename: str = "document.docx") -> bytes:
    """Return .docx bytes for .docx, ZIP-misnamed .doc, or LibreOffice-converted .doc."""
    if not file_bytes:
        raise ValueError(f"Uploaded Word file '{filename}' is empty.")

    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".docx" or (ext == ".doc" and _is_zip_docx_bytes(file_bytes)):
        return file_bytes
    if ext == ".doc":
        return _convert_doc_to_docx_bytes(file_bytes)
    raise ValueError(f"Unsupported Word extension '{ext}' for '{filename}'.")


def load_word_text_and_images(
    file_bytes: bytes, filename: str = "document.docx"
) -> Tuple[str, List[Tuple[str, bytes]]]:
    """
    Load Word content as (body_text, embedded_images).

    body_text may be empty when the order is a pasted screenshot image.
    """
    docx_bytes = normalize_word_to_docx_bytes(file_bytes, filename)
    text = extract_text_from_docx_bytes(docx_bytes)
    images = extract_images_from_docx_bytes(docx_bytes)
    return text, images


def extract_text_from_word_bytes(file_bytes: bytes, filename: str = "document.docx") -> str:
    """
    Extract text-only from .docx/.doc (no OCR).

    Prefer load_word_text_and_images() when image fallback OCR is needed.
    """
    text, _images = load_word_text_and_images(file_bytes, filename)
    if not text:
        raise ValueError(
            f"No readable text found in '{filename}'. "
            "If the Word file is scanned/image-only, export to PDF or an image and re-upload."
        )
    return text
